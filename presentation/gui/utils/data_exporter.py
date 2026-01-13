# data_exporter.py
"""
M√≥dulo para exportar datos de Treeview a diferentes formatos.
"""
import keyword
import platform
import subprocess
import tempfile
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from typing import Optional, List, Any
import os
from datetime import datetime
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch



HAS_EXCEL = True
HAS_WORD = True
HAS_PDF = True

# Agregar al inicio de data_exporter.py, despu√©s de las importaciones
class DataHierarchyTransformer:
    """Transformador de datos planos a estructura jer√°rquica"""
    
    @staticmethod
    def detect_key_columns(headers):
        """
        Detecta las columnas clave en los headers
        
        Returns:
            dict: √çndices de columnas importantes
        """
        indices = {
            'department': None,
            'employee': None,
            'amount': None
        }
        
        # Mapeo de posibles nombres para cada columna
        department_keywords = ['departamento', 'depto', 'unidad', 'area', 'direccion', 'gerencia', 'Departamento']
        employee_keywords = ['solicitante', 'empleado', 'nombre', 'fullname', 'colaborador', 'trabajador', 'persona', 'Solicitante']
        amount_keywords = ['gasto', 'monto', 'total', 'importe', 'cantidad', 'costo', 'precio', 'valor', 'saldo', 'monto liquidado', 'Monto Liquidado', 'Monto']
        
        for i, header in enumerate(headers):
            header_lower = str(header).lower()
            
            # Buscar departamento
            if indices['department'] is None:
                for keyword in department_keywords:
                    if keyword in header_lower:
                        indices['department'] = i
                        break
            
            # Buscar empleado/solicitante
            if indices['employee'] is None:
                for keyword in employee_keywords:
                    if keyword in header_lower:
                        indices['employee'] = i
                        break
            
            # Buscar monto/gasto
            if indices['amount'] is None:
                for keyword in amount_keywords:
                    if keyword in header_lower:
                        indices['amount'] = i
                        break
        
        return indices
    
    @staticmethod
    def transform_to_hierarchical(headers, data):
        """
        Transforma datos planos a estructura jer√°rquica
        
        Returns:
            dict: {
                'headers': headers originales,
                'hierarchical_data': datos transformados con jerarqu√≠a,
                'summary': resumen por departamento,
                'total_general': total general
            }
        """
        if not headers or not data:
            return None
        
        # Detectar columnas clave
        indices = DataHierarchyTransformer.detect_key_columns(headers)
        
        # Si no hay columna de departamento, no podemos hacer jerarqu√≠a
        if indices['department'] is None:
            return None
        
        # Agrupar por departamento
        departments = {}
        for row in data:
            dept = str(row[indices['department']]) if indices['department'] < len(row) else "SIN DEPARTAMENTO"
            if dept not in departments:
                departments[dept] = {
                    'rows': [],
                    'subtotal': 0
                }
            
            departments[dept]['rows'].append(row)
            
            # Calcular subtotal si hay columna de monto
            if indices['amount'] is not None and indices['amount'] < len(row):
                try:
                    amount_str = str(row[indices['amount']]).replace('$', '').replace(',', '').replace(' ', '')
                    if amount_str.replace('.', '', 1).isdigit():
                        amount = float(amount_str)
                        departments[dept]['subtotal'] += amount
                except (ValueError, AttributeError):
                    pass
        
        # Construir estructura jer√°rquica
        hierarchical_data = []
        total_general = 0
        summary = {}
        
        # Ordenar departamentos alfab√©ticamente
        sorted_depts = sorted(departments.keys())
        
        for dept in sorted_depts:
            dept_info = departments[dept]
            
            # 1. Fila de encabezado del departamento
            dept_header = [''] * len(headers)
            dept_header[indices['department']] = f"üìä DEPARTAMENTO: {dept.upper()}"
            
            # Agregar subtotal si existe
            if indices['amount'] is not None and dept_info['subtotal'] > 0:
                # Buscar la columna de monto para poner el subtotal
                subtotal_text = f"SUBTOTAL: ${dept_info['subtotal']:,.2f}"
                dept_header[indices['amount']] = subtotal_text
            
            hierarchical_data.append({
                'type': 'department_header',
                'data': dept_header,
                'department': dept,
                'subtotal': dept_info['subtotal']
            })
            
            summary[dept] = dept_info['subtotal']
            total_general += dept_info['subtotal']
            
            # 2. Filas de solicitantes (con numeraci√≥n)
            for idx, row in enumerate(dept_info['rows'], 1):
                formatted_row = list(row)
                
                # Numerar solicitantes
                if indices['employee'] is not None and indices['employee'] < len(formatted_row):
                    employee_name = formatted_row[indices['employee']]
                    formatted_row[indices['employee']] = f"{idx}. {employee_name}"
                
                hierarchical_data.append({
                    'type': 'employee_row',
                    'data': formatted_row,
                    'department': dept,
                    'employee_number': idx
                })
            
            # 3. Separador entre departamentos
            hierarchical_data.append({
                'type': 'separator',
                'data': [''] * len(headers),
                'department': dept
            })
        
        # 4. Fila de total general
        if indices['amount'] is not None and total_general > 0:
            total_row = [''] * len(headers)
            total_row[indices['department']] = "‚úÖ TOTAL GENERAL"
            total_row[indices['amount']] = f"${total_general:,.2f}"
            
            hierarchical_data.append({
                'type': 'total_row',
                'data': total_row,
                'total': total_general
            })
        
        return {
            'headers': headers,
            'hierarchical_data': hierarchical_data,
            'summary': summary,
            'total_general': total_general,
            'indices': indices
        }


class TreeviewExporter:
    """Exportador gen√©rico para Treeview de tkinter"""
    
    
    @staticmethod
    def get_treeview_data(tree: ttk.Treeview, hierarchical=True):
        """
        Extrae todos los datos del Treeview con transformaci√≥n jer√°rquica autom√°tica
        
        Args:
            hierarchical: Si True, intenta transformar a estructura jer√°rquica
        
        Returns:
            tuple: (encabezados, datos, estructura_jerarquica)
        """
        # Obtener encabezados
        columns = tree['columns']
        headers = []
        for col in columns:
            header = tree.heading(col)['text']
            headers.append(header)
        
        # Obtener datos
        data = []
        for item in tree.get_children():
            values = tree.item(item)['values']
            data.append(values)
        
        # Aplicar transformaci√≥n jer√°rquica si est√° habilitada y hay datos
        hierarchical_structure = None
        if hierarchical and headers and data:
            hierarchical_structure = DataHierarchyTransformer.transform_to_hierarchical(headers, data)
        
        return headers, data, hierarchical_structure
    
    @staticmethod
    def export_to_excel(tree: ttk.Treeview, title: str, filename: str = None) -> Optional[str]:
        """Exporta Treeview a Excel con tablas separadas por departamento"""
        if not HAS_EXCEL:
            messagebox.showerror("Error", "openpyxl no est√° instalado. Inst√°lelo con: pip install openpyxl")
            return None
        
        if not filename:
            filename = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")],
                title="Guardar como Excel"
            )
            
        if not filename:
            return None
        
        try:
            # Obtener datos con transformaci√≥n jer√°rquica autom√°tica
            headers, _, hierarchical_structure = TreeviewExporter.get_treeview_data(tree, hierarchical=True)
            
            if not headers:
                messagebox.showwarning("Sin datos", "No hay datos para exportar")
                return None
            
            # Funci√≥n para abreviar encabezados espec√≠ficos
            def abreviar_encabezado(header):
                header_str = str(header)
                abreviaciones = {
                    'Desayunos': 'D',
                    'Almuerzos': 'A', 
                    'Cenas': 'C',
                    'Alojamientos': 'H',
                }
                for key, value in abreviaciones.items():
                    if key in header_str:
                        return value
                return header_str
            
            # Crear workbook
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Reporte"
            
            # Escribir t√≠tulo principal
            ws.merge_cells(f'A1:{get_column_letter(len(headers))}1')
            title_cell = ws['A1']
            title_cell.value = title
            title_cell.font = Font(size=14, bold=True, color="2c3e50")
            title_cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # Subt√≠tulo informativo si hay jerarqu√≠a
            if hierarchical_structure:
                ws.merge_cells(f'A2:{get_column_letter(len(headers))}2')
                subtitle = "üìä REPORTE POR DEPARTAMENTOS"
                subtitle_cell = ws['A2']
                subtitle_cell.value = subtitle
                subtitle_cell.font = Font(size=11, bold=True, color="27ae60")
                subtitle_cell.alignment = Alignment(horizontal='center')
                subtitle_cell.fill = PatternFill(start_color="e8f8f5", end_color="e8f8f5", fill_type="solid")
            
            # Fecha de exportaci√≥n
            ws.merge_cells(f'A3:{get_column_letter(len(headers))}3')
            date_cell = ws['A3']
            date_cell.value = f"Exportado: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
            date_cell.font = Font(size=9, color="666666")
            date_cell.alignment = Alignment(horizontal='center')
            
            # Espacio
            current_row = 5
            
            if hierarchical_structure:
                # PROCESAR POR DEPARTAMENTOS
                departments_data = {}
                current_dept = None
                
                for item in hierarchical_structure['hierarchical_data']:
                    if item['type'] == 'department_header':
                        current_dept = item['department']
                        departments_data[current_dept] = {
                            'header': item,
                            'employees': [],
                            'subtotal': item.get('subtotal', 0)
                        }
                    elif item['type'] == 'employee_row' and current_dept:
                        departments_data[current_dept]['employees'].append(item)
                
                # CREAR UNA SECCI√ìN POR CADA DEPARTAMENTO
                for dept_name, dept_info in departments_data.items():
                    # T√≠tulo del departamento con subtotal
                    dept_cell = ws.cell(row=current_row, column=1)
                    dept_text = f"üìä {dept_name}"
                    if dept_info['subtotal'] > 0:
                        dept_text += f" - Subtotal: ${dept_info['subtotal']:,.2f}"
                    
                    dept_cell.value = dept_text
                    dept_cell.font = Font(bold=True, color="2c3e50", size=12)
                    dept_cell.fill = PatternFill(start_color="e8f4f8", end_color="e8f4f8", fill_type="solid")
                    
                    # Fusionar celdas para el t√≠tulo del departamento
                    ws.merge_cells(
                        start_row=current_row,
                        start_column=1,
                        end_row=current_row,
                        end_column=len(headers)
                    )
                    
                    current_row += 1
                    
                    # Crear tabla para este departamento
                    num_employees = len(dept_info['employees'])
                    if num_employees > 0:
                        # Identificar √≠ndices de columnas a excluir
                        indices_a_excluir = []
                        
                        # 1. Columna de departamento (siempre)
                        if hierarchical_structure['indices'].get('department') is not None:
                            indices_a_excluir.append(hierarchical_structure['indices']['department'])
                        
                        # 2. Columna "N¬∞ Anticipo" y "N¬∞ Liquidaci√≥n"
                        for i, header in enumerate(headers):
                            header_lower = str(header).lower()
                            if any(keyword in header_lower for keyword in ['anticipo', 'n¬∞ anticipo', 'n anticipo', 'n¬∫ anticipo', 'numero anticipo',
                                                                        'no. anticipo', 'liquidaci√≥n', 'n¬∞ liquidaci√≥n', 'n liquidaci√≥n', 
                                                                        'n¬∫ liquidaci√≥n', 'numero liquidaci√≥n', 'no. liquidaci√≥n']):
                                if i not in indices_a_excluir:
                                    indices_a_excluir.append(i)
                        
                        # 3. Columna "Estado"
                        for i, header in enumerate(headers):
                            header_lower = str(header).lower()
                            if any(keyword in header_lower for keyword in ['estado', 'status', 'situacion']):
                                if i not in indices_a_excluir:
                                    indices_a_excluir.append(i)
                        
                        indices_a_excluir.sort()
                        
                        # ENCABEZADOS DE LA TABLA (excluyendo columnas no deseadas)
                        col_idx_destino = 1
                        
                        # Lista para mapear √≠ndices de columna a headers originales
                        included_headers = []
                        
                        for i, header in enumerate(headers):
                            if i in indices_a_excluir:
                                continue
                            
                            included_headers.append(header)  # Guardar header para referencia
                            header_cell = ws.cell(row=current_row, column=col_idx_destino)
                            header_abreviado = abreviar_encabezado(header)
                            header_cell.value = str(header_abreviado)
                            header_cell.font = Font(bold=True, color="FFFFFF")
                            header_cell.fill = PatternFill(start_color="2c3e50", end_color="2c3e50", fill_type="solid")
                            header_cell.alignment = Alignment(horizontal='center', vertical='center')
                            
                            # Ajustar ancho de columna
                            col_letter = get_column_letter(col_idx_destino)
                            ws.column_dimensions[col_letter].width = max(len(str(header_abreviado)) + 4, 12)
                            
                            col_idx_destino += 1
                        
                        current_row += 1
                        
                        # DATOS DE LOS EMPLEADOS
                        for idx, emp_item in enumerate(dept_info['employees']):
                            row_data = emp_item['data']
                            col_idx_destino = 1
                            
                            for i, cell_data in enumerate(row_data):
                                if i in indices_a_excluir:
                                    continue
                                
                                cell = ws.cell(row=current_row, column=col_idx_destino, value=cell_data)
                                
                                # Obtener el header original para esta columna
                                original_header = included_headers[col_idx_destino - 1]
                                
                                # Formatear nombre del empleado con numeraci√≥n
                                if i == hierarchical_structure['indices'].get('employee'):
                                    employee_name = cell_data
                                    clean_name = str(employee_name).replace('   ‚îú‚îÄ‚îÄ ', '').replace('‚îú‚îÄ‚îÄ ', '')
                                    cell.value = f"{idx + 1}. {clean_name}"
                                    cell.alignment = Alignment(horizontal='left')
                                
                                # Formatear montos (excluyendo columnas D, A, C, H)
                                cell_str = str(cell_data)
                                
                                # Verificar si es una columna de cantidad (D, A, C, H)
                                is_quantity_column = False
                                if original_header and any(keyword in str(original_header).lower() for keyword in ['desayunos', 'almuerzos', 'cenas', 'alojamientos']):
                                    is_quantity_column = True
                                
                                # Para montos (excepto columnas de cantidad)
                                if cell_str.startswith('$') or (cell_str.replace('.', '', 1).replace(',', '').isdigit() and not is_quantity_column):
                                    cell.alignment = Alignment(horizontal='right')
                                    cell.number_format = '"$"#,##0.00'
                                # Para columnas de cantidad (D, A, C, H)
                                elif is_quantity_column and cell_str.replace('.', '', 1).isdigit():
                                    # Mostrar como n√∫mero entero
                                    try:
                                        # Convertir a entero si es decimal
                                        if '.' in cell_str:
                                            cell.value = int(float(cell_str))
                                        else:
                                            cell.value = int(cell_str)
                                        cell.alignment = Alignment(horizontal='right')
                                        cell.number_format = '0'  # Formato entero sin decimales
                                    except (ValueError, TypeError):
                                        pass
                                
                                # Fondo alternado para filas
                                if idx % 2 == 0:
                                    cell.fill = PatternFill(start_color="F8F9FA", end_color="F8F9FA", fill_type="solid")
                                
                                col_idx_destino += 1
                            
                            current_row += 1
                    
                    # Espacio entre departamentos
                    current_row += 1
                
                # TABLA DE RESUMEN FINAL
                summary_row = current_row + 1
                summary_title = ws.cell(row=summary_row, column=1, value="üìä RESUMEN DE DEPARTAMENTOS")
                summary_title.font = Font(bold=True, color="2c3e50", size=12)
                ws.merge_cells(
                    start_row=summary_row,
                    start_column=1,
                    end_row=summary_row,
                    end_column=2
                )
                
                summary_row += 1
                
                # Encabezados del resumen
                ws.cell(row=summary_row, column=1, value="DEPARTAMENTO").font = Font(bold=True)
                ws.cell(row=summary_row, column=2, value="SUBTOTAL").font = Font(bold=True)
                
                summary_row += 1
                
                # Datos del resumen
                total_general = 0
                for dept_name, dept_info in departments_data.items():
                    ws.cell(row=summary_row, column=1, value=dept_name)
                    ws.cell(row=summary_row, column=2, value=f"${dept_info['subtotal']:,.2f}")
                    ws.cell(row=summary_row, column=2).number_format = '"$"#,##0.00'
                    total_general += dept_info['subtotal']
                    summary_row += 1
                
                # Fila de total general
                total_row = summary_row
                ws.cell(row=total_row, column=1, value="TOTAL GENERAL").font = Font(bold=True, color="FFFFFF")
                ws.cell(row=total_row, column=2, value=f"${total_general:,.2f}").font = Font(bold=True, color="FFFFFF")
                ws.cell(row=total_row, column=2).number_format = '"$"#,##0.00'
                
                # Formato para celdas de total
                for col in [1, 2]:
                    cell = ws.cell(row=total_row, column=col)
                    cell.fill = PatternFill(start_color="27ae60", end_color="27ae60", fill_type="solid")
            else:
                # FALLBACK: Tabla plana (sin jerarqu√≠a)
                _, flat_data, _ = TreeviewExporter.get_treeview_data(tree, hierarchical=False)
                
                # Escribir encabezados
                for col_idx, header in enumerate(headers, start=1):
                    cell = ws.cell(row=current_row, column=col_idx, value=header)
                    cell.font = Font(bold=True, color="FFFFFF")
                    cell.fill = PatternFill(start_color="2c3e50", end_color="2c3e50", fill_type="solid")
                    cell.alignment = Alignment(horizontal='center', vertical='center')
                    
                    col_letter = get_column_letter(col_idx)
                    ws.column_dimensions[col_letter].width = max(len(str(header)) + 4, 15)
                
                current_row += 1
                
                # Escribir datos
                for row_idx, row_data in enumerate(flat_data):
                    for col_idx, cell_data in enumerate(row_data, start=1):
                        cell = ws.cell(row=current_row, column=col_idx, value=cell_data)
                        
                        if isinstance(cell_data, str) and cell_data.startswith('$'):
                            cell.alignment = Alignment(horizontal='right')
                        elif isinstance(cell_data, (int, float)):
                            cell.alignment = Alignment(horizontal='right')
                            cell.number_format = '"$"#,##0.00'
                        
                        if row_idx % 2 == 0:
                            cell.fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
                    
                    current_row += 1
            
            # Agregar bordes a todas las celdas con datos
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            
            # Aplicar bordes din√°micamente
            max_row = ws.max_row
            max_col = ws.max_column
            
            for row in ws.iter_rows(min_row=1, max_row=max_row, min_col=1, max_col=max_col):
                for cell in row:
                    if cell.value:
                        cell.border = thin_border
            
            # Pie de p√°gina
            footer_row = max_row + 2
            ws.cell(row=footer_row, column=1, 
                value=f"Sistema de Gesti√≥n de Dietas VIAJEX ‚Ä¢ Exportado el: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
            ws.cell(row=footer_row, column=1).font = Font(italic=True, size=9, color="666666")
            

            for col_idx in range(1, ws.max_column + 1):
                max_length = 0
                col_letter = get_column_letter(col_idx)
                
                # Verificar si esta columna tiene dimensiones definidas
                if col_letter in ws.column_dimensions:
                    # Iterar solo sobre filas que existen
                    for row in range(1, ws.max_row + 1):
                        cell = ws.cell(row=row, column=col_idx)
                        
                        # Verificar que no sea una celda fusionada
                        try:
                            # Si es MergedCell, saltarla
                            if hasattr(cell, 'is_merged') and cell.is_merged:
                                continue
                        except:
                            pass
                        
                        # Calcular longitud del contenido
                        if cell.value:
                            try:
                                cell_length = len(str(cell.value))
                                if cell_length > max_length:
                                    max_length = cell_length
                            except:
                                pass
                    
                    # Ajustar ancho (m√≠nimo 10, m√°ximo 50)
                    adjusted_width = min(max(max_length + 2, 10), 50)
                    ws.column_dimensions[col_letter].width = adjusted_width
            
            wb.save(filename)
            
            # Mensaje de √©xito
            if hierarchical_structure:
                message = (
                    f"‚úÖ EXCEL :\n\n"
                    f"üìÇ {filename}\n\n"
                    f"‚Ä¢ Tablas por departamento: {len(departments_data)}\n"
                    f"‚Ä¢ Total General: ${total_general:,.2f}\n"
                    f"‚Ä¢ Incluye resumen detallado"
                )
            else:
                message = f"‚úÖ Excel exportado exitosamente:\n\n{filename}"
            
            messagebox.showinfo("Exportaci√≥n exitosa", message)
            return filename
            
        except Exception as e:
            messagebox.showerror("‚ùå Error", f"No se pudo exportar a Excel:\n\n{str(e)}")
            import traceback
            traceback.print_exc()
            return None        
        
    @staticmethod
    def export_to_word(tree: ttk.Treeview, title: str, filename: str = None) -> Optional[str]:
        """Exporta Treeview a Word con tablas separadas por departamento"""
        if not HAS_WORD:
            messagebox.showerror("Error", "python-docx no est√° instalado. Inst√°lelo con: pip install python-docx")
            return None
        
        if not filename:
            filename = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx"), ("All files", "*.*")],
                title="Guardar como Word (Tablas separadas)"
            )
            
        if not filename:
            return None
        
        try:
            
            # Obtener datos con transformaci√≥n jer√°rquica autom√°tica
            headers, _, hierarchical_structure = TreeviewExporter.get_treeview_data(tree, hierarchical=True)
            
            if not headers:
                messagebox.showwarning("Sin datos", "No hay datos para exportar")
                return None
            
            # Crear documento
            doc = Document()

            def abreviar_encabezado(header):
                """Abrevia encabezados espec√≠ficos a su inicial"""
                header_str = str(header)
                
                # Diccionario de abreviaciones
                abreviaciones = {
                    'Desayunos': 'D',
                    'Almuerzos': 'A', 
                    'Cenas': 'C',
                    'Alojamientos': 'H',  # H para Hospedaje
                }
                
                # Verificar si el encabezado est√° en el diccionario
                for key, value in abreviaciones.items():
                    if key in header_str:
                        return value
                
                # Si no es un encabezado espec√≠fico, devolver el original
                return header_str
            
            # Configurar p√°gina
            section = doc.sections[0]
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            
            # T√≠tulo principal
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.font.size = Pt(14)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(44, 62, 80)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Subt√≠tulo informativo si hay jerarqu√≠a
            if hierarchical_structure:
                subtitle_para = doc.add_paragraph()
                subtitle_run = subtitle_para.add_run("üìä REPORTE POR DEPARTAMENTOS")
                subtitle_run.font.size = Pt(11)
                subtitle_run.font.bold = True
                subtitle_run.font.color.rgb = RGBColor(39, 174, 96)
                subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Fecha de exportaci√≥n
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"Exportado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            date_run.font.size = Pt(9)
            date_run.font.color.rgb = RGBColor(102, 102, 102)
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Espacio
            
            if hierarchical_structure:
                # PROCESAR POR DEPARTAMENTOS - TABLAS SEPARADAS
                # Primero, agrupar los datos por departamento
                departments_data = {}
                
                current_dept = None
                for item in hierarchical_structure['hierarchical_data']:
                    if item['type'] == 'department_header':
                        current_dept = item['department']
                        departments_data[current_dept] = {
                            'header': item,
                            'employees': [],
                            'subtotal': item.get('subtotal', 0)
                        }
                    elif item['type'] == 'employee_row' and current_dept:
                        departments_data[current_dept]['employees'].append(item)
                    
                
                # CREAR UNA TABLA POR CADA DEPARTAMENTO
                for dept_name, dept_info in departments_data.items():
                    # T√≠tulo del departamento (como p√°rrafo, no como fila de tabla)
                    dept_title_para = doc.add_paragraph()
                    
                    # Formatear el texto del departamento con subtotal
                    dept_header_text = dept_info['header']['data'][hierarchical_structure['indices']['department']]
                    if dept_info['subtotal'] > 0:
                        subtotal_text = f" - Subtotal: ${dept_info['subtotal']:,.2f}"
                    else:
                        subtotal_text = ""
                    
                    dept_title_run = dept_title_para.add_run(f"üìä {dept_header_text}{subtotal_text}")
                    dept_title_run.font.size = Pt(12)
                    dept_title_run.font.bold = True
                    dept_title_run.font.color.rgb = RGBColor(44, 62, 80)
    

                    # Crear tabla para este departamento
                    num_employees = len(dept_info['employees'])
                    if num_employees > 0:
                        # Identificar √≠ndices de columnas a excluir
                        indices_a_excluir = []
                        
                        # 1. Columna de departamento (siempre)
                        if hierarchical_structure['indices'].get('department') is not None:
                            indices_a_excluir.append(hierarchical_structure['indices']['department'])
                        
                        # 2. Columna "N¬∞ Anticipo" - buscar por varios nombres posibles
                        for i, header in enumerate(headers):
                            header_lower = str(header).lower()
                            if any(keyword in header_lower for keyword in ['anticipo', 'n¬∞ anticipo', 'n anticipo', 'n¬∫ anticipo', 'numero anticipo',
                                                                        'no. anticipo', 'liquidaci√≥n', 'n¬∞ liquidaci√≥n', 'n liquidaci√≥n', 
                                                                        'n¬∫ liquidaci√≥n', 'numero liquidaci√≥n', 'no. liquidaci√≥n']):
                                if i not in indices_a_excluir:
                                    indices_a_excluir.append(i)
                        
                        # 3. Columna "Estado" - buscar por varios nombres posibles
                        for i, header in enumerate(headers):
                            header_lower = str(header).lower()
                            if any(keyword in header_lower for keyword in ['estado', 'status', 'situacion']):
                                if i not in indices_a_excluir:
                                    indices_a_excluir.append(i)
                        
                        # Ordenar los √≠ndices a excluir
                        indices_a_excluir.sort()
                        
                        # Calcular n√∫mero de columnas finales
                        num_columnas_finales = len(headers) - len(indices_a_excluir)
                        
                        # Crear tabla con las columnas restantes
                        table = doc.add_table(rows=num_employees + 1, cols=num_columnas_finales)
                        table.style = 'Table Grid'
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        table.autofit = False
                        
                        # Configurar anchos de columna proporcionales
                        for i in range(num_columnas_finales):
                            table.columns[i].width = Inches(1.8)  # Un poco m√°s ancho porque hay menos columnas
                        
                        # ENCABEZADOS DE LA TABLA (excluyendo las columnas no deseadas)
                        header_cells = table.rows[0].cells
                        col_idx_destino = 0
                        
                        for i, header in enumerate(headers):
                            # Saltar las columnas que est√°n en la lista de exclusi√≥n
                            if i in indices_a_excluir:
                                continue
                            
                            # Abreviar encabezados espec√≠ficos
                            header_abreviado = abreviar_encabezado(header)
                            header_cells[col_idx_destino].text = str(header_abreviado)
                            
                            header_cells[col_idx_destino].paragraphs[0].runs[0].font.bold = True
                            header_cells[col_idx_destino].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            header_cells[col_idx_destino].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                            
                            # Fondo azul oscuro
                            tc = header_cells[col_idx_destino]._tc
                            tcPr = tc.get_or_add_tcPr()
                            shading = OxmlElement('w:shd')
                            shading.set(qn('w:fill'), '2c3e50')
                            tcPr.append(shading)
                            
                            col_idx_destino += 1
                        
                        # DATOS DE LOS EMPLEADOS (excluyendo las mismas columnas)
                        for idx, emp_item in enumerate(dept_info['employees']):
                            row_cells = table.rows[idx + 1].cells
                            row_data = emp_item['data']
                            
                            col_idx_destino = 0
                            for i, cell_data in enumerate(row_data):
                                # Saltar las columnas que est√°n en la lista de exclusi√≥n
                                if i in indices_a_excluir:
                                    continue
                                    
                                # Formatear el nombre del empleado con numeraci√≥n
                                if i == hierarchical_structure['indices'].get('employee'):
                                    employee_name = cell_data
                                    # Quitar cualquier prefijo existente y agregar numeraci√≥n
                                    clean_name = str(employee_name).replace('   ‚îú‚îÄ‚îÄ ', '').replace('‚îú‚îÄ‚îÄ ', '')
                                    row_cells[col_idx_destino].text = f"{idx + 1}. {clean_name}"
                                else:
                                    row_cells[col_idx_destino].text = str(cell_data)
                                
                                # Formato para montos
                                cell_str = str(cell_data)
                                if cell_str.startswith('$') or cell_str.replace('.', '', 1).replace(',', '').isdigit():
                                    row_cells[col_idx_destino].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                
                                # Fondo alternado para filas
                                if idx % 2 == 0:
                                    tc = row_cells[col_idx_destino]._tc
                                    tcPr = tc.get_or_add_tcPr()
                                    shading = OxmlElement('w:shd')
                                    shading.set(qn('w:fill'), 'f8f9fa')
                                    tcPr.append(shading)
                                
                                col_idx_destino += 1
                        
                        # Ajustar altura de filas
                        for row in table.rows:
                            tr = row._tr
                            trPr = tr.get_or_add_trPr()
                            trHeight = OxmlElement('w:trHeight')
                            trHeight.set(qn('w:val'), "350")
                            trPr.append(trHeight)

                    # Espacio entre departamentos
                    doc.add_paragraph()
                
                # TABLA DE RESUMEN FINAL (departamentos con subtotales)
                summary_title = doc.add_paragraph()
                summary_title_run = summary_title.add_run("üìä RESUMEN DE DEPARTAMENTOS")
                summary_title_run.font.size = Pt(12)
                summary_title_run.font.bold = True
                summary_title_run.font.color.rgb = RGBColor(44, 62, 80)
                
                # Crear tabla de resumen (2 columnas: departamento, subtotal)
                summary_table = doc.add_table(rows=len(departments_data) + 2, cols=2)
                summary_table.style = 'Table Grid'
                summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Encabezados del resumen
                summary_header = summary_table.rows[0].cells
                summary_header[0].text = "DEPARTAMENTO"
                summary_header[1].text = "SUBTOTAL"
                
                for cell in summary_header:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                    
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), '2c3e50')
                    tcPr.append(shading)
                
                # Datos del resumen
                row_idx = 1
                total_general = 0
                
                for dept_name, dept_info in departments_data.items():
                    row_cells = summary_table.rows[row_idx].cells
                    row_cells[0].text = dept_name
                    row_cells[1].text = f"${dept_info['subtotal']:,.2f}"
                    row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    total_general += dept_info['subtotal']
                    row_idx += 1
                
                # Fila de total general
                total_cells = summary_table.rows[row_idx].cells
                total_cells[0].text = "TOTAL GENERAL"
                total_cells[1].text = f"${total_general:,.2f}"
                
                for cell in total_cells:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                    
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), '27ae60')
                    tcPr.append(shading)
                
                total_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                # FALLBACK: Tabla plana (sin jerarqu√≠a)
                _, flat_data, _ = TreeviewExporter.get_treeview_data(tree, hierarchical=False)
                
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Configurar anchos
                for i in range(len(headers)):
                    table.columns[i].width = Inches(1.5)
                
                # Encabezados
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    header_cells[i].text = str(header)
                    header_cells[i].paragraphs[0].runs[0].font.bold = True
                    header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                    
                    tc = header_cells[i]._tc
                    tcPr = tc.get_or_add_tcPr()
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), '2c3e50')
                    tcPr.append(shading)
                
                # Datos
                for row_idx, row_data in enumerate(flat_data):
                    row_cells = table.add_row().cells
                    
                    for i, cell_data in enumerate(row_data):
                        row_cells[i].text = str(cell_data)
                        
                        # Fondo alternado
                        if row_idx % 2 == 0:
                            tc = row_cells[i]._tc
                            tcPr = tc.get_or_add_tcPr()
                            shading = OxmlElement('w:shd')
                            shading.set(qn('w:fill'), 'f8f9fa')
                            tcPr.append(shading)
            
            # Ajustar altura de filas en todas las tablas
            for table in doc.tables:
                for row in table.rows:
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = OxmlElement('w:trHeight')
                    trHeight.set(qn('w:val'), "350")
                    trPr.append(trHeight)
            
            # Pie de documento
            doc.add_paragraph()
            footer_para = doc.add_paragraph()
            
            footer_text = "Sistema de Gesti√≥n de Dietas VIAJEX"
            if hierarchical_structure:
                footer_text += f" ‚Ä¢ {len(departments_data)} departamentos procesados"
            
            footer_run = footer_para.add_run(footer_text)
            footer_run.font.size = Pt(8)
            footer_run.italic = True
            footer_run.font.color.rgb = RGBColor(102, 102, 102)
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Guardar documento
            doc.save(filename)
            
            # Mensaje de √©xito
            if hierarchical_structure:
                message = (
                    f"‚úÖ WORD con tablas separadas exportado:\n\n"
                    f"üìÇ {filename}\n\n"
                    f"‚Ä¢ Tablas por departamento: {len(departments_data)}\n"
                    f"‚Ä¢ Total General: ${total_general:,.2f}\n"
                    f"‚Ä¢ Incluye resumen detallado"
                )
            else:
                message = f"‚úÖ Word exportado exitosamente:\n\n{filename}"
            
            messagebox.showinfo("Exportaci√≥n exitosa", message)
            return filename
            
        except Exception as e:
            messagebox.showerror("‚ùå Error", f"No se pudo exportar a Word:\n\n{str(e)}")
            import traceback
            traceback.print_exc()
            return None        
        
    @staticmethod
    def export_to_pdf(tree: ttk.Treeview, title: str, filename: str = None) -> Optional[str]:
        """Exporta Treeview a PDF con tablas separadas por departamento"""
        if not HAS_PDF:
            messagebox.showerror("Error", "reportlab no est√° instalado. Inst√°lelo con: pip install reportlab")
            return None
        
        if not filename:
            filename = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")],
                title="Guardar como PDF (Tablas separadas)"
            )
            
        if not filename:
            return None
        
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib import colors
            from reportlab.lib.units import inch
            
            # Obtener datos con transformaci√≥n jer√°rquica autom√°tica
            headers, _, hierarchical_structure = TreeviewExporter.get_treeview_data(tree, hierarchical=True)
            
            if not headers:
                messagebox.showwarning("Sin datos", "No hay datos para exportar")
                return None
            
            # Funci√≥n para abreviar encabezados espec√≠ficos
            def abreviar_encabezado(header):
                header_str = str(header)
                abreviaciones = {
                    'Desayunos': 'D',
                    'Almuerzos': 'A', 
                    'Cenas': 'C',
                    'Alojamientos': 'H',
                }
                for key, value in abreviaciones.items():
                    if key in header_str:
                        return value
                return header_str
            
            # Crear documento PDF con m√°rgenes optimizados
            doc = SimpleDocTemplate(
                filename, 
                pagesize=A4,
                topMargin=0.5*inch, 
                bottomMargin=0.5*inch,
                leftMargin=0.4*inch,  # Reducido para m√°s espacio
                rightMargin=0.4*inch
            )
            elements = []
            
            # Estilos
            styles = getSampleStyleSheet()
            
            # T√≠tulo principal
            title_style = ParagraphStyle(
                'MainTitle',
                parent=styles['Heading1'],
                fontSize=14,
                textColor=colors.HexColor('#2c3e50'),
                spaceAfter=8,
                alignment=1,
                fontName='Helvetica-Bold'
            )
            
            main_title = Paragraph(title, title_style)
            elements.append(main_title)
            
            # Subt√≠tulo informativo si hay jerarqu√≠a
            if hierarchical_structure:
                subtitle_style = ParagraphStyle(
                    'Subtitle',
                    parent=styles['Normal'],
                    fontSize=11,
                    textColor=colors.HexColor('#27ae60'),
                    alignment=1,
                    spaceAfter=10
                )
                
                subtitle = Paragraph("üìä REPORTE POR DEPARTAMENTOS", subtitle_style)
                elements.append(subtitle)
            
            # Fecha de exportaci√≥n
            date_style = ParagraphStyle(
                'DateStyle',
                parent=styles['Normal'],
                fontSize=9,
                textColor=colors.HexColor('#666666'),
                alignment=1,
                spaceAfter=12
            )
            
            date_para = Paragraph(f"Exportado: {datetime.now().strftime('%d/%m/%Y %H:%M')}", date_style)
            elements.append(date_para)
            
            elements.append(Spacer(1, 0.15*inch))
            
            if hierarchical_structure:
                # PROCESAR POR DEPARTAMENTOS
                departments_data = {}
                current_dept = None
                
                for item in hierarchical_structure['hierarchical_data']:
                    if item['type'] == 'department_header':
                        current_dept = item['department']
                        departments_data[current_dept] = {
                            'header': item,
                            'employees': [],
                            'subtotal': item.get('subtotal', 0)
                        }
                    elif item['type'] == 'employee_row' and current_dept:
                        departments_data[current_dept]['employees'].append(item)
                
                total_general = 0
                
                # CREAR UNA SECCI√ìN POR CADA DEPARTAMENTO
                for dept_idx, (dept_name, dept_info) in enumerate(departments_data.items()):
                    # T√≠tulo del departamento
                    dept_title_style = ParagraphStyle(
                        'DeptTitle',
                        parent=styles['Heading2'],
                        fontSize=11,
                        textColor=colors.HexColor('#2c3e50'),
                        spaceAfter=8,
                        leftIndent=0,
                        fontName='Helvetica-Bold',
                        backColor=colors.HexColor('#e8f4f8'),
                        borderPadding=(6, 6, 6, 6)
                    )
                    
                    dept_text = f"üìä {dept_name}"
                    if dept_info['subtotal'] > 0:
                        dept_text += f" - Subtotal: ${dept_info['subtotal']:,.2f}"
                        total_general += dept_info['subtotal']
                    
                    dept_title = Paragraph(dept_text, dept_title_style)
                    elements.append(dept_title)
                    
                    # Crear tabla para este departamento
                    num_employees = len(dept_info['employees'])
                    if num_employees > 0:
                        # Identificar √≠ndices de columnas a excluir
                        indices_a_excluir = []
                        
                        # 1. Columna de departamento (siempre)
                        if hierarchical_structure['indices'].get('department') is not None:
                            indices_a_excluir.append(hierarchical_structure['indices']['department'])
                        
                        # 2. Columna "N¬∞ Anticipo" y "N¬∞ Liquidaci√≥n"
                        for i, header in enumerate(headers):
                            header_lower = str(header).lower()
                            if any(keyword in header_lower for keyword in ['anticipo', 'n¬∞ anticipo', 'n anticipo', 'n¬∫ anticipo', 'numero anticipo',
                                                                        'no. anticipo', 'liquidaci√≥n', 'n¬∞ liquidaci√≥n', 'n liquidaci√≥n', 
                                                                        'n¬∫ liquidaci√≥n', 'numero liquidaci√≥n', 'no. liquidaci√≥n']):
                                if i not in indices_a_excluir:
                                    indices_a_excluir.append(i)
                        
                        # 3. Columna "Estado"
                        for i, header in enumerate(headers):
                            header_lower = str(header).lower()
                            if any(keyword in header_lower for keyword in ['estado', 'status', 'situacion']):
                                if i not in indices_a_excluir:
                                    indices_a_excluir.append(i)
                        
                        indices_a_excluir.sort()
                        
                        # Construir encabezados de la tabla
                        table_headers = []
                        for i, header in enumerate(headers):
                            if i in indices_a_excluir:
                                continue
                            header_abreviado = abreviar_encabezado(header)
                            table_headers.append(header_abreviado)
                        
                        # Construir datos de la tabla
                        table_data = [table_headers]
                        
                        for idx, emp_item in enumerate(dept_info['employees']):
                            row_data = emp_item['data']
                            table_row = []
                            
                            for i, cell_data in enumerate(row_data):
                                if i in indices_a_excluir:
                                    continue
                                
                                # Formatear nombre del empleado con numeraci√≥n usando Paragraph para ajuste de texto
                                if i == hierarchical_structure['indices'].get('employee'):
                                    employee_name = cell_data
                                    clean_name = str(employee_name).replace('   ‚îú‚îÄ‚îÄ ', '').replace('‚îú‚îÄ‚îÄ ', '')
                                    # Crear Paragraph con estilo que permita wrap de texto
                                    employee_style = ParagraphStyle(
                                        'EmployeeStyle',
                                        parent=styles['Normal'],
                                        fontSize=8,
                                        wordWrap='CJK',  # Permite wrap de texto
                                        leading=10,  # Espacio entre l√≠neas
                                    )
                                    table_row.append(Paragraph(f"{idx + 1}. {clean_name}", employee_style))
                                else:
                                    table_row.append(str(cell_data) if cell_data is not None else "")
                            
                            table_data.append(table_row)
                        
                        # Crear tabla
                        if table_data and len(table_data[0]) > 0:
                            num_cols = len(table_data[0])
                            
                            # Calcular anchos de columna proporcionales
                            col_widths = []
                            total_width = doc.width
                            
                            # Distribuci√≥n inteligente de anchos
                            if num_cols <= 4:
                                # Poco columnas: distribuir equitativamente
                                col_width = total_width / num_cols
                                col_widths = [col_width] * num_cols
                            else:
                                # Muchas columnas: dar m√°s espacio a la primera (empleado)
                                col_widths = [total_width * 0.35]  # 35% para empleado
                                # Resto dividido entre las otras columnas
                                remaining_width = total_width * 0.65
                                other_cols_width = remaining_width / (num_cols - 1)
                                col_widths.extend([other_cols_width] * (num_cols - 1))
                            
                            table = Table(table_data, colWidths=col_widths, repeatRows=1)
                            
                            # ESTILO DE LA TABLA MEJORADO
                            table_style = TableStyle([
                                # Encabezados
                                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                ('FONTSIZE', (0, 0), (-1, 0), 9),
                                ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
                                ('TOPPADDING', (0, 0), (-1, 0), 6),
                                
                                # Bordes finos
                                ('GRID', (0, 0), (-1, -1), 0.25, colors.HexColor("#dddddd")),
                                
                                # Filas alternadas con mejor contraste
                                ('ROWBACKGROUNDS', (0, 1), (-1, -1), 
                                [colors.white, colors.HexColor("#f5f7fa")]),
                                
                                # Alineaci√≥n y padding mejorado
                                ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
                                ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # Alinear arriba para mejor ajuste
                                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                                ('FONTSIZE', (0, 1), (-1, -1), 8),
                                
                                # Padding interno para evitar superposici√≥n
                                ('LEFTPADDING', (0, 1), (-1, -1), 8),
                                ('RIGHTPADDING', (0, 1), (-1, -1), 8),
                                ('TOPPADDING', (0, 1), (-1, -1), 4),
                                ('BOTTOMPADDING', (0, 1), (-1, -1), 4),
                                
                                # Ajuste especial para la columna del empleado (m√°s espacio)
                                ('LEFTPADDING', (0, 1), (0, -1), 12),
                                ('RIGHTPADDING', (0, 1), (0, -1), 10),
                            ])
                            
                            # Alinear montos a la derecha
                            for col_idx, header in enumerate(table_headers):
                                if any(keyword in str(header).lower() for keyword in ['gasto', 'monto', 'total', 'precio', 'costo', '$', 'saldo']):
                                    table_style.add('ALIGN', (col_idx, 1), (col_idx, -1), 'RIGHT')
                            
                            table.setStyle(table_style)
                            elements.append(table)
                            elements.append(Spacer(1, 0.25*inch))
                    
                    # Espacio entre departamentos
                    if dept_idx < len(departments_data) - 1:
                        elements.append(Spacer(1, 0.15*inch))
                
                # Salto de p√°gina para el resumen
                elements.append(PageBreak())
                
                # Tabla de resumen final
                summary_title_style = ParagraphStyle(
                    'SummaryTitle',
                    parent=styles['Heading2'],
                    fontSize=12,
                    textColor=colors.HexColor('#2c3e50'),
                    spaceAfter=12,
                    alignment=1
                )
                
                summary_title = Paragraph("üìä RESUMEN DE DEPARTAMENTOS", summary_title_style)
                elements.append(summary_title)
                
                # Crear tabla de resumen
                summary_data = [["DEPARTAMENTO", "SUBTOTAL"]]
                
                for dept_name, dept_info in departments_data.items():
                    summary_data.append([dept_name, f"${dept_info['subtotal']:,.2f}"])
                
                summary_data.append(["TOTAL GENERAL", f"${total_general:,.2f}"])
                
                summary_table = Table(summary_data, colWidths=[doc.width * 0.7, doc.width * 0.3])
                
                # Estilo mejorado para la tabla de resumen
                summary_style = TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#dddddd")),
                    
                    ('ALIGN', (1, 1), (1, -2), 'RIGHT'),
                    ('FONTNAME', (1, -1), (1, -1), 'Helvetica-Bold'),
                    ('BACKGROUND', (1, -1), (1, -1), colors.HexColor("#27ae60")),
                    ('TEXTCOLOR', (1, -1), (1, -1), colors.white),
                    
                    # Padding mejorado
                    ('LEFTPADDING', (0, 1), (-1, -1), 10),
                    ('RIGHTPADDING', (0, 1), (-1, -1), 10),
                    ('TOPPADDING', (0, 1), (-1, -1), 6),
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
                ])
                
                # Tambi√©n aplicar estilo a la primera columna del total
                summary_style.add('BACKGROUND', (0, -1), (0, -1), colors.HexColor("#27ae60"))
                summary_style.add('TEXTCOLOR', (0, -1), (0, -1), colors.white)
                
                summary_table.setStyle(summary_style)
                elements.append(summary_table)
            else:
                # FALLBACK: Tabla plana (sin jerarqu√≠a)
                _, flat_data, _ = TreeviewExporter.get_treeview_data(tree, hierarchical=False)
                
                # Construir tabla completa
                table_data = [headers] + flat_data
                
                num_cols = len(headers)
                if num_cols > 0:
                    col_widths = [doc.width / num_cols] * num_cols
                    table = Table(table_data, colWidths=col_widths, repeatRows=1)
                    
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#dddddd")),
                        
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), 
                        [colors.white, colors.HexColor("#f5f7fa")]),
                        
                        # Padding para evitar superposici√≥n
                        ('LEFTPADDING', (0, 1), (-1, -1), 8),
                        ('RIGHTPADDING', (0, 1), (-1, -1), 8),
                    ]))
                    
                    elements.append(table)
            
            # Pie de p√°gina mejorado
            elements.append(Spacer(1, 0.3*inch))
            
            footer_style = ParagraphStyle(
                'Footer',
                parent=styles['Italic'],
                fontSize=8,
                textColor=colors.HexColor('#666666'),
                alignment=1,
                spaceBefore=10
            )
            
            footer_text = "Sistema de Gesti√≥n de Dietas VIAJEX"
            if hierarchical_structure:
                footer_text += f" ‚Ä¢ {len(departments_data)} departamentos procesados"
            
            footer = Paragraph(footer_text, footer_style)
            elements.append(footer)
            
            # Construir documento
            doc.build(elements)
            
            # Mensaje de √©xito
            if hierarchical_structure:
                message = (
                    f"‚úÖ PDF con tablas separadas exportado:\n\n"
                    f"üìÇ {filename}\n\n"
                    f"‚Ä¢ Tablas por departamento: {len(departments_data)}\n"
                    f"‚Ä¢ Total General: ${total_general:,.2f}\n"
                    f"‚Ä¢ Incluye resumen detallado"
                )
            else:
                message = f"‚úÖ PDF exportado exitosamente:\n\n{filename}"
            
            messagebox.showinfo("Exportaci√≥n exitosa", message)
            return filename
            
        except Exception as e:
            messagebox.showerror("‚ùå Error", f"No se pudo exportar a PDF:\n\n{str(e)}")
            import traceback
            traceback.print_exc()
            return None    
                
    @staticmethod
    def show_dependency_help():
        """Muestra ayuda para instalar dependencias"""
        help_text = """
            Para exportar en todos los formatos, instale las siguientes bibliotecas:

            Excel: pip install openpyxl
            Word:  pip install python-docx
            PDF:   pip install reportlab

            O instale todas con:
            pip install openpyxl python-docx reportlab
        """
        messagebox.showinfo("Dependencias requeridas", help_text)

    @staticmethod
    def print_directly(tree: ttk.Treeview, title: str) -> bool:
        """Imprime directamente sin mostrar vista previa"""
        try:
            # Crear un PDF temporal con el MISMO formato que export_to_pdf
            temp_dir = tempfile.gettempdir()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            temp_filename = os.path.join(temp_dir, f"impresion_{timestamp}.pdf")
                   
            # Obtener datos con transformaci√≥n jer√°rquica autom√°tica
            headers, _, hierarchical_structure = TreeviewExporter.get_treeview_data(tree, hierarchical=True)
            
            if not headers:
                messagebox.showwarning("Sin datos", "No hay datos para imprimir")
                return False
            
            # Funci√≥n para abreviar encabezados espec√≠ficos (ID√âNTICA a export_to_pdf)
            def abreviar_encabezado(header):
                header_str = str(header)
                abreviaciones = {
                    'Desayunos': 'D',
                    'Almuerzos': 'A', 
                    'Cenas': 'C',
                    'Alojamientos': 'H',
                }
                for key, value in abreviaciones.items():
                    if key in header_str:
                        return value
                return header_str
            
            # Crear documento PDF con m√°rgenes optimizados (ID√âNTICO a export_to_pdf)
            doc = SimpleDocTemplate(
                temp_filename, 
                pagesize=A4,
                topMargin=0.5*inch, 
                bottomMargin=0.5*inch,
                leftMargin=0.4*inch,
                rightMargin=0.4*inch
            )
            elements = []
            
            # Estilos (ID√âNTICOS a export_to_pdf)
            styles = getSampleStyleSheet()
            
            # T√≠tulo principal (ID√âNTICO a export_to_pdf)
            title_style = ParagraphStyle(
                'MainTitle',
                parent=styles['Heading1'],
                fontSize=14,
                textColor=colors.HexColor('#2c3e50'),
                spaceAfter=8,
                alignment=1,
                fontName='Helvetica-Bold'
            )
            
            main_title = Paragraph(title, title_style)
            elements.append(main_title)
            
            # Subt√≠tulo informativo si hay jerarqu√≠a (ID√âNTICO a export_to_pdf)
            if hierarchical_structure:
                subtitle_style = ParagraphStyle(
                    'Subtitle',
                    parent=styles['Normal'],
                    fontSize=11,
                    textColor=colors.HexColor('#27ae60'),
                    alignment=1,
                    spaceAfter=10
                )
                
                subtitle = Paragraph("üìä REPORTE POR DEPARTAMENTOS", subtitle_style)
                elements.append(subtitle)
            
            # Fecha de exportaci√≥n (ID√âNTICO a export_to_pdf)
            date_style = ParagraphStyle(
                'DateStyle',
                parent=styles['Normal'],
                fontSize=9,
                textColor=colors.HexColor('#666666'),
                alignment=1,
                spaceAfter=12
            )
            
            date_para = Paragraph(f"Impreso: {datetime.now().strftime('%d/%m/%Y %H:%M')}", date_style)
            elements.append(date_para)
            
            elements.append(Spacer(1, 0.15*inch))
            
            if hierarchical_structure:
                # PROCESAR POR DEPARTAMENTOS 
                departments_data = {}
                current_dept = None
                
                for item in hierarchical_structure['hierarchical_data']:
                    if item['type'] == 'department_header':
                        current_dept = item['department']
                        departments_data[current_dept] = {
                            'header': item,
                            'employees': [],
                            'subtotal': item.get('subtotal', 0)
                        }
                    elif item['type'] == 'employee_row' and current_dept:
                        departments_data[current_dept]['employees'].append(item)
                
                total_general = 0
                
                # CREAR UNA SECCI√ìN POR CADA DEPARTAMENTO 
                for dept_idx, (dept_name, dept_info) in enumerate(departments_data.items()):
                    # T√≠tulo del departamento
                    dept_title_style = ParagraphStyle(
                        'DeptTitle',
                        parent=styles['Heading2'],
                        fontSize=11,
                        textColor=colors.HexColor('#2c3e50'),
                        spaceAfter=8,
                        leftIndent=0,
                        fontName='Helvetica-Bold',
                        backColor=colors.HexColor('#e8f4f8'),
                        borderPadding=(6, 6, 6, 6)
                    )
                    
                    dept_text = f"üìä {dept_name}"
                    if dept_info['subtotal'] > 0:
                        dept_text += f" - Subtotal: ${dept_info['subtotal']:,.2f}"
                        total_general += dept_info['subtotal']
                    
                    dept_title = Paragraph(dept_text, dept_title_style)
                    elements.append(dept_title)
                    
                    # Crear tabla para este departamento 
                    num_employees = len(dept_info['employees'])
                    if num_employees > 0:
                        # Identificar √≠ndices de columnas a excluir 
                        indices_a_excluir = []
                        
                        # 1. Columna de departamento (siempre)
                        if hierarchical_structure['indices'].get('department') is not None:
                            indices_a_excluir.append(hierarchical_structure['indices']['department'])
                        
                        # 2. Columna "N¬∞ Anticipo" y "N¬∞ Liquidaci√≥n"
                        for i, header in enumerate(headers):
                            header_lower = str(header).lower()
                            if any(keyword in header_lower for keyword in ['anticipo', 'n¬∞ anticipo', 'n anticipo', 'n¬∫ anticipo', 'numero anticipo',
                                                                        'no. anticipo', 'liquidaci√≥n', 'n¬∞ liquidaci√≥n', 'n liquidaci√≥n', 
                                                                        'n¬∫ liquidaci√≥n', 'numero liquidaci√≥n', 'no. liquidaci√≥n']):
                                if i not in indices_a_excluir:
                                    indices_a_excluir.append(i)
                        
                        # 3. Columna "Estado"
                        for i, header in enumerate(headers):
                            header_lower = str(header).lower()
                            if any(keyword in header_lower for keyword in ['estado', 'status', 'situacion']):
                                if i not in indices_a_excluir:
                                    indices_a_excluir.append(i)
                        
                        indices_a_excluir.sort()
                        
                        # Construir encabezados de la tabla (ID√âNTICO a export_to_pdf)
                        table_headers = []
                        for i, header in enumerate(headers):
                            if i in indices_a_excluir:
                                continue
                            header_abreviado = abreviar_encabezado(header)
                            table_headers.append(header_abreviado)
                        
                        # Construir datos de la tabla (ID√âNTICO a export_to_pdf)
                        table_data = [table_headers]
                        
                        for idx, emp_item in enumerate(dept_info['employees']):
                            row_data = emp_item['data']
                            table_row = []
                            
                            for i, cell_data in enumerate(row_data):
                                if i in indices_a_excluir:
                                    continue
                                
                                # Formatear nombre del empleado con numeraci√≥n usando Paragraph para ajuste de texto
                                if i == hierarchical_structure['indices'].get('employee'):
                                    employee_name = cell_data
                                    clean_name = str(employee_name).replace('   ‚îú‚îÄ‚îÄ ', '').replace('‚îú‚îÄ‚îÄ ', '')
                                    # Crear Paragraph con estilo que permita wrap de texto
                                    employee_style = ParagraphStyle(
                                        'EmployeeStyle',
                                        parent=styles['Normal'],
                                        fontSize=8,
                                        wordWrap='CJK',  # Permite wrap de texto
                                        leading=10,  # Espacio entre l√≠neas
                                    )
                                    table_row.append(Paragraph(f"{idx + 1}. {clean_name}", employee_style))
                                else:
                                    table_row.append(str(cell_data) if cell_data is not None else "")
                            
                            table_data.append(table_row)
                        
                        # Crear tabla 
                        if table_data and len(table_data[0]) > 0:
                            num_cols = len(table_data[0])
                            
                            # Calcular anchos de columna proporcionales 
                            col_widths = []
                            total_width = doc.width
                            
                            # Distribuci√≥n inteligente de anchos 
                            if num_cols <= 4:
                                # Poco columnas: distribuir equitativamente
                                col_width = total_width / num_cols
                                col_widths = [col_width] * num_cols
                            else:
                                # Muchas columnas: dar m√°s espacio a la primera (empleado)
                                col_widths = [total_width * 0.35]  # 35% para empleado
                                # Resto dividido entre las otras columnas
                                remaining_width = total_width * 0.65
                                other_cols_width = remaining_width / (num_cols - 1)
                                col_widths.extend([other_cols_width] * (num_cols - 1))
                            
                            table = Table(table_data, colWidths=col_widths, repeatRows=1)
                            
                            # ESTILO DE LA TABLA MEJORADO (ID√âNTICO a export_to_pdf)
                            table_style = TableStyle([
                                # Encabezados
                                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
                                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                ('FONTSIZE', (0, 0), (-1, 0), 9),
                                ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
                                ('TOPPADDING', (0, 0), (-1, 0), 6),
                                
                                # Bordes finos
                                ('GRID', (0, 0), (-1, -1), 0.25, colors.HexColor("#dddddd")),
                                
                                # Filas alternadas con mejor contraste
                                ('ROWBACKGROUNDS', (0, 1), (-1, -1), 
                                [colors.white, colors.HexColor("#f5f7fa")]),
                                
                                # Alineaci√≥n y padding mejorado
                                ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
                                ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # Alinear arriba para mejor ajuste
                                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                                ('FONTSIZE', (0, 1), (-1, -1), 8),
                                
                                # Padding interno para evitar superposici√≥n
                                ('LEFTPADDING', (0, 1), (-1, -1), 8),
                                ('RIGHTPADDING', (0, 1), (-1, -1), 8),
                                ('TOPPADDING', (0, 1), (-1, -1), 4),
                                ('BOTTOMPADDING', (0, 1), (-1, -1), 4),
                                
                                # Ajuste especial para la columna del empleado (m√°s espacio)
                                ('LEFTPADDING', (0, 1), (0, -1), 12),
                                ('RIGHTPADDING', (0, 1), (0, -1), 10),
                            ])
                            
                            # Alinear montos a la derecha (ID√âNTICO a export_to_pdf)
                            for col_idx, header in enumerate(table_headers):
                                if any(keyword in str(header).lower() for keyword in ['gasto', 'monto', 'total', 'precio', 'costo', '$', 'saldo']):
                                    table_style.add('ALIGN', (col_idx, 1), (col_idx, -1), 'RIGHT')
                            
                            table.setStyle(table_style)
                            elements.append(table)
                            elements.append(Spacer(1, 0.25*inch))
                    
                    # Espacio entre departamentos (ID√âNTICO a export_to_pdf)
                    if dept_idx < len(departments_data) - 1:
                        elements.append(Spacer(1, 0.15*inch))
                
                # Salto de p√°gina para el resumen (ID√âNTICO a export_to_pdf)
                elements.append(PageBreak())
                
                # Tabla de resumen final (ID√âNTICO a export_to_pdf)
                summary_title_style = ParagraphStyle(
                    'SummaryTitle',
                    parent=styles['Heading2'],
                    fontSize=12,
                    textColor=colors.HexColor('#2c3e50'),
                    spaceAfter=12,
                    alignment=1
                )
                
                summary_title = Paragraph("üìä RESUMEN DE DEPARTAMENTOS", summary_title_style)
                elements.append(summary_title)
                
                # Crear tabla de resumen (ID√âNTICO a export_to_pdf)
                summary_data = [["DEPARTAMENTO", "SUBTOTAL"]]
                
                for dept_name, dept_info in departments_data.items():
                    summary_data.append([dept_name, f"${dept_info['subtotal']:,.2f}"])
                
                summary_data.append(["TOTAL GENERAL", f"${total_general:,.2f}"])
                
                summary_table = Table(summary_data, colWidths=[doc.width * 0.7, doc.width * 0.3])
                
                # Estilo mejorado para la tabla de resumen (ID√âNTICO a export_to_pdf)
                summary_style = TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#dddddd")),
                    
                    ('ALIGN', (1, 1), (1, -2), 'RIGHT'),
                    ('FONTNAME', (1, -1), (1, -1), 'Helvetica-Bold'),
                    ('BACKGROUND', (1, -1), (1, -1), colors.HexColor("#27ae60")),
                    ('TEXTCOLOR', (1, -1), (1, -1), colors.white),
                    
                    # Padding mejorado
                    ('LEFTPADDING', (0, 1), (-1, -1), 10),
                    ('RIGHTPADDING', (0, 1), (-1, -1), 10),
                    ('TOPPADDING', (0, 1), (-1, -1), 6),
                    ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
                ])
                
                # Tambi√©n aplicar estilo a la primera columna del total (ID√âNTICO a export_to_pdf)
                summary_style.add('BACKGROUND', (0, -1), (0, -1), colors.HexColor("#27ae60"))
                summary_style.add('TEXTCOLOR', (0, -1), (0, -1), colors.white)
                
                summary_table.setStyle(summary_style)
                elements.append(summary_table)
            else:
                # FALLBACK: Tabla plana (sin jerarqu√≠a) - ID√âNTICO a export_to_pdf
                _, flat_data, _ = TreeviewExporter.get_treeview_data(tree, hierarchical=False)
                
                # Construir tabla completa
                table_data = [headers] + flat_data
                
                num_cols = len(headers)
                if num_cols > 0:
                    col_widths = [doc.width / num_cols] * num_cols
                    table = Table(table_data, colWidths=col_widths, repeatRows=1)
                    
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#dddddd")),
                        
                        ('ROWBACKGROUNDS', (0, 1), (-1, -1), 
                        [colors.white, colors.HexColor("#f5f7fa")]),
                        
                        # Padding para evitar superposici√≥n
                        ('LEFTPADDING', (0, 1), (-1, -1), 8),
                        ('RIGHTPADDING', (0, 1), (-1, -1), 8),
                    ]))
                    
                    elements.append(table)
            
            # Pie de p√°gina mejorado (ID√âNTICO a export_to_pdf)
            elements.append(Spacer(1, 0.3*inch))
            
            footer_style = ParagraphStyle(
                'Footer',
                parent=styles['Italic'],
                fontSize=8,
                textColor=colors.HexColor('#666666'),
                alignment=1,
                spaceBefore=10
            )
            
            footer_text = "Sistema de Gesti√≥n de Dietas VIAJEX"
            if hierarchical_structure:
                footer_text += f" ‚Ä¢ {len(departments_data)} departamentos procesados"
            
            footer = Paragraph(footer_text, footer_style)
            elements.append(footer)
            
            # Construir documento 
            doc.build(elements)
            
            # Ahora abrir el PDF para imprimir 
            system = platform.system()
            
            if system == "Windows":
                try:
                    
                    import time
                    time.sleep(1)  # Esperar un segundo para que se abra el PDF
                    
                    messagebox.showinfo(
                        "üìÑ Imprimir documento",
                        "Se ha abierto el documento PDF en su visor predeterminado.\n\n"
                        "Para imprimir:\n"
                        "1. Presione Ctrl+P en el visor de PDF\n"
                        "2. Seleccione su impresora\n"
                        "3. Ajuste las configuraciones si es necesario\n"
                        "4. Haga clic en 'Imprimir'\n\n"
                        f"Archivo: {temp_filename}\n"
                        f"Este archivo se eliminar√° autom√°ticamente."
                    )
                    
                    os.startfile(temp_filename)
                except Exception as e:
                    # Si no se puede abrir, mostrar el archivo en el explorador
                    messagebox.showinfo(
                        "üìÑ Localizar archivo para imprimir",
                        f"No se pudo abrir el PDF autom√°ticamente.\n\n"
                        f"Por favor, abra manualmente el archivo:\n\n"
                        f"{temp_filename}\n\n"
                        f"Y luego impr√≠malo con Ctrl+P."
                    )
                    # Abrir el explorador en la carpeta
                    os.startfile(temp_dir)
            
            # Programar eliminaci√≥n del archivo temporal (despu√©s de 60 segundos)
            def delete_temp_file():
                try:
                    if os.path.exists(temp_filename):
                        os.remove(temp_filename)
                except:
                    pass
            
            import threading
            timer = threading.Timer(60.0, delete_temp_file)
            timer.start()
            
            return True
            
        except Exception as e:
            messagebox.showerror("‚ùå Error de impresi√≥n", 
                            f"No se pudo preparar la impresi√≥n:\n\n{str(e)}")
            import traceback
            traceback.print_exc()
            return False

    @staticmethod
    def export_to_excel_full_columns(tree: ttk.Treeview, title: str, filename: str = None) -> Optional[str]:
        """Exporta Treeview a Excel con todas las columnas del reporte"""
        if not HAS_EXCEL:
            messagebox.showerror("Error", "openpyxl no est√° instalado. Inst√°lelo con: pip install openpyxl")
            return None
        
        if not filename:
            filename = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")],
                title="Guardar como Excel (Todas las columnas)"
            )
            
        if not filename:
            return None
        
        try:
            # Obtener datos del treeview
            columns = tree['columns']
            headers = []
            for col in columns:
                header = tree.heading(col)['text']
                headers.append(header)
            
            data = []
            for item in tree.get_children():
                values = tree.item(item)['values']
                data.append(values)
            
            if not headers or not data:
                messagebox.showwarning("Sin datos", "No hay datos para exportar")
                return None
            
            # Detectar √≠ndices de columnas importantes
            def detect_columns():
                indices = {
                    'departamento': None,
                    'solicitante': None,
                    'monto_efec': None,
                    'monto_card': None,
                    'gasto_efec': None,
                    'gasto_card': None,
                    'estado': None
                }
                
                for i, header in enumerate(headers):
                    header_str = str(header).lower()
                    
                    if indices['departamento'] is None:
                        if any(keyword in header_str for keyword in ['departamento', 'depto', 'unidad', 'area']):
                            indices['departamento'] = i
                    
                    if indices['solicitante'] is None:
                        if any(keyword in header_str for keyword in ['solicitante', 'empleado', 'nombre', 'fullname']):
                            indices['solicitante'] = i
                    
                    if indices['monto_efec'] is None:
                        if any(keyword in header_str for keyword in ['s.e', 'efec', 'efectivo solicitado', 'monto efec']):
                            indices['monto_efec'] = i
                    
                    if indices['monto_card'] is None:
                        if any(keyword in header_str for keyword in ['s.t', 'tarjeta solicitado', 'monto card', 'tarjeta']):
                            indices['monto_card'] = i
                    
                    if indices['gasto_efec'] is None:
                        if any(keyword in header_str for keyword in ['g.e', 'gasto efec', 'gasto efectivo']):
                            indices['gasto_efec'] = i
                    
                    if indices['gasto_card'] is None:
                        if any(keyword in header_str for keyword in ['g.t', 'gasto card', 'gasto tarjeta']):
                            indices['gasto_card'] = i
                    
                    if indices['estado'] is None:
                        if any(keyword in header_str for keyword in ['estado', 'status', 'situacion']):
                            indices['estado'] = i
                
                return indices
            
            indices = detect_columns()
            
            # Crear workbook
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Reporte Completo"
            
            # Escribir t√≠tulo principal
            ws.merge_cells(f'A1:{get_column_letter(len(headers))}1')
            title_cell = ws['A1']
            title_cell.value = f"{title} - Todas las columnas"
            title_cell.font = Font(size=14, bold=True, color="2c3e50")
            title_cell.alignment = Alignment(horizontal='center', vertical='center')
            
            # Fecha de exportaci√≥n
            ws.merge_cells(f'A2:{get_column_letter(len(headers))}2')
            date_cell = ws['A2']
            date_cell.value = f"Exportado: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
            date_cell.font = Font(size=9, color="666666")
            date_cell.alignment = Alignment(horizontal='center')
            
            # Comenzar en la fila 4
            current_row = 4
            
            # Agrupar datos por departamento si se detect√≥ columna de departamento
            if indices['departamento'] is not None:
                # Agrupar por departamento
                departments = {}
                for row in data:
                    dept = str(row[indices['departamento']]) if indices['departamento'] < len(row) else "Sin Departamento"
                    if dept not in departments:
                        departments[dept] = []
                    departments[dept].append(row)
                
                # Ordenar departamentos
                sorted_depts = sorted(departments.keys())
                total_general_efec = 0
                total_general_card = 0
                
                for dept in sorted_depts:
                    dept_data = departments[dept]
                    
                    # T√≠tulo del departamento
                    dept_cell = ws.cell(row=current_row, column=1)
                    dept_cell.value = f"üìä DEPARTAMENTO: {dept.upper()}"
                    dept_cell.font = Font(bold=True, color="2c3e50", size=12)
                    dept_cell.fill = PatternFill(start_color="e8f4f8", end_color="e8f4f8", fill_type="solid")
                    
                    ws.merge_cells(
                        start_row=current_row,
                        start_column=1,
                        end_row=current_row,
                        end_column=len(headers)
                    )
                    
                    current_row += 1
                    
                    # Encabezados de tabla
                    for col_idx, header in enumerate(headers, start=1):
                        header_cell = ws.cell(row=current_row, column=col_idx)
                        header_cell.value = header
                        header_cell.font = Font(bold=True, color="FFFFFF")
                        header_cell.fill = PatternFill(start_color="2c3e50", end_color="2c3e50", fill_type="solid")
                        header_cell.alignment = Alignment(horizontal='center', vertical='center')
                        
                        col_letter = get_column_letter(col_idx)
                        ws.column_dimensions[col_letter].width = max(len(str(header)) + 2, 12)
                    
                    current_row += 1
                    
                    # Variables para subtotales del departamento
                    dept_subtotal_efec = 0
                    dept_subtotal_card = 0
                    dept_gasto_efec = 0
                    dept_gasto_card = 0
                    
                    # Datos del departamento
                    for idx, row in enumerate(dept_data):
                        for col_idx, cell_data in enumerate(row, start=1):
                            cell = ws.cell(row=current_row, column=col_idx, value=cell_data)
                            
                            # Formatear seg√∫n tipo de dato
                            cell_str = str(cell_data) if cell_data is not None else ""
                            
                            # Convertir montos a n√∫meros (quitar $ y comas)
                            is_monto_column = False
                            if col_idx - 1 in [indices['monto_efec'], indices['monto_card'], 
                                              indices['gasto_efec'], indices['gasto_card']]:
                                is_monto_column = True
                            
                            if is_monto_column and cell_str:
                                # Quitar s√≠mbolos de moneda y caracteres no num√©ricos
                                clean_str = cell_str.replace('$', '').replace(',', '').replace(' ', '').strip()
                                
                                # Verificar si es un n√∫mero
                                if clean_str.replace('.', '', 1).isdigit():
                                    try:
                                        num_value = float(clean_str)
                                        cell.value = num_value
                                        
                                        # Formato num√©rico con 2 decimales
                                        cell.number_format = '0.00'
                                        cell.alignment = Alignment(horizontal='right')
                                        
                                        # Acumular subtotales
                                        if col_idx - 1 == indices['monto_efec']:
                                            dept_subtotal_efec += num_value
                                            total_general_efec += num_value
                                        elif col_idx - 1 == indices['monto_card']:
                                            dept_subtotal_card += num_value
                                            total_general_card += num_value
                                        elif col_idx - 1 == indices['gasto_efec']:
                                            dept_gasto_efec += num_value
                                        elif col_idx - 1 == indices['gasto_card']:
                                            dept_gasto_card += num_value
                                            
                                    except (ValueError, TypeError):
                                        # Si no se puede convertir, mantener el string original
                                        cell.value = cell_str
                                else:
                                    cell.value = cell_str
                            
                            # Numerar solicitantes
                            elif col_idx - 1 == indices['solicitante']:
                                if cell_str:
                                    # Quitar prefijos existentes y agregar numeraci√≥n
                                    clean_name = cell_str.replace('   ‚îú‚îÄ‚îÄ ', '').replace('‚îú‚îÄ‚îÄ ', '')
                                    cell.value = f"{idx + 1}. {clean_name}"
                            
                            # Formatear fechas
                            elif any(keyword in headers[col_idx-1].lower() for keyword in ['fecha', 'date']):
                                cell.alignment = Alignment(horizontal='center')
                            
                            # Formatear estado
                            elif col_idx - 1 == indices['estado']:
                                cell.alignment = Alignment(horizontal='center')
                        
                        # Fondo alternado para filas
                        if idx % 2 == 0:
                            for col_idx in range(1, len(headers) + 1):
                                ws.cell(row=current_row, column=col_idx).fill = PatternFill(
                                    start_color="F8F9FA", end_color="F8F9FA", fill_type="solid"
                                )
                        
                        current_row += 1
                    
                    # Fila de subtotales del departamento
                    subtotal_row = current_row
                    
                    # Crear fila de subtotal
                    for col_idx in range(1, len(headers) + 1):
                        cell = ws.cell(row=subtotal_row, column=col_idx)
                        
                        # Poner subtotales en sus columnas respectivas
                        if col_idx - 1 == indices['monto_efec']:
                            cell.value = dept_subtotal_efec
                            cell.number_format = '0.00'
                            cell.font = Font(bold=True, color="2c3e50")
                            cell.fill = PatternFill(start_color="FFF3CD", end_color="FFF3CD", fill_type="solid")
                        elif col_idx - 1 == indices['monto_card']:
                            cell.value = dept_subtotal_card
                            cell.number_format = '0.00'
                            cell.font = Font(bold=True, color="2c3e50")
                            cell.fill = PatternFill(start_color="FFF3CD", end_color="FFF3CD", fill_type="solid")
                        elif col_idx - 1 == indices['departamento']:
                            cell.value = f"Subtotal {dept}"
                            cell.font = Font(bold=True, italic=True)
                            cell.fill = PatternFill(start_color="FFF3CD", end_color="FFF3CD", fill_type="solid")
                    
                    current_row += 2  # Espacio entre departamentos
                
                                # --- RESUMEN POR DEPARTAMENTO (NUEVO) ---
                summary_row = current_row + 1
                
                # T√≠tulo del resumen por departamento
                ws.merge_cells(f'A{summary_row}:{get_column_letter(len(headers))}{summary_row}')
                summary_title = ws.cell(row=summary_row, column=1)
                summary_title.value = "üìä RESUMEN POR DEPARTAMENTOS"
                summary_title.font = Font(bold=True, color="2c3e50", size=12)
                summary_title.alignment = Alignment(horizontal='center')
                summary_title.fill = PatternFill(start_color="E8F8F5", end_color="E8F8F5", fill_type="solid")
                
                summary_row += 1
                
                # Encabezados del resumen por departamento
                dept_summary_headers = [
                    "DEPARTAMENTO",
                    "Total Efectivo Anticipado",
                    "Total Tarjeta Anticipado",
                    "Total Efectivo Gastado",
                    "Total Tarjeta Gastado",
                    "Reembolso Efectivo",
                    "Reembolso Tarjeta",
                    "Gasto Total",
                    "Subtotal Departamento"
                ]
                
                for col_idx, header in enumerate(dept_summary_headers, start=1):
                    header_cell = ws.cell(row=summary_row, column=col_idx, value=header)
                    header_cell.font = Font(bold=True, color="FFFFFF")
                    header_cell.fill = PatternFill(start_color="2c3e50", end_color="2c3e50", fill_type="solid")
                    header_cell.alignment = Alignment(horizontal='center', vertical='center')
                    
                    # Ajustar ancho de columna
                    col_letter = get_column_letter(col_idx)
                    ws.column_dimensions[col_letter].width = max(len(header) + 2, 15)
                
                summary_row += 1
                
                # Variables para el resumen general
                general_efec_anticipado = 0
                general_card_anticipado = 0
                general_efec_gastado = 0
                general_card_gastado = 0
                
                # Crear lista para almacenar resumen por departamento
                dept_summary_data = []
                
                # Procesar cada departamento para el resumen
                for dept in sorted_depts:
                    dept_data = departments[dept]
                    
                    # Inicializar acumuladores para este departamento
                    dept_efec_anticipado = 0
                    dept_card_anticipado = 0
                    dept_efec_gastado = 0
                    dept_card_gastado = 0
                    
                    # Calcular montos por departamento
                    for row in dept_data:
                        # Efectivo solicitado (S.E)
                        if indices['monto_efec'] is not None and indices['monto_efec'] < len(row):
                            cell_str = str(row[indices['monto_efec']]) if row[indices['monto_efec']] is not None else ""
                            if cell_str:
                                clean_str = cell_str.replace('$', '').replace(',', '').replace(' ', '').strip()
                                if clean_str.replace('.', '', 1).isdigit():
                                    dept_efec_anticipado += float(clean_str)
                                    general_efec_anticipado += float(clean_str)
                        
                        # Tarjeta solicitado (S.T)
                        if indices['monto_card'] is not None and indices['monto_card'] < len(row):
                            cell_str = str(row[indices['monto_card']]) if row[indices['monto_card']] is not None else ""
                            if cell_str:
                                clean_str = cell_str.replace('$', '').replace(',', '').replace(' ', '').strip()
                                if clean_str.replace('.', '', 1).isdigit():
                                    dept_card_anticipado += float(clean_str)
                                    general_card_anticipado += float(clean_str)
                        
                        # Gasto efectivo (G.E)
                        if indices['gasto_efec'] is not None and indices['gasto_efec'] < len(row):
                            cell_str = str(row[indices['gasto_efec']]) if row[indices['gasto_efec']] is not None else ""
                            if cell_str:
                                clean_str = cell_str.replace('$', '').replace(',', '').replace(' ', '').strip()
                                if clean_str.replace('.', '', 1).isdigit():
                                    dept_efec_gastado += float(clean_str)
                                    general_efec_gastado += float(clean_str)
                        
                        # Gasto tarjeta (G.T)
                        if indices['gasto_card'] is not None and indices['gasto_card'] < len(row):
                            cell_str = str(row[indices['gasto_card']]) if row[indices['gasto_card']] is not None else ""
                            if cell_str:
                                clean_str = cell_str.replace('$', '').replace(',', '').replace(' ', '').strip()
                                if clean_str.replace('.', '', 1).isdigit():
                                    dept_card_gastado += float(clean_str)
                                    general_card_gastado += float(clean_str)
                    
                    # Calcular valores derivados
                    reembolso_efec = dept_efec_anticipado - dept_efec_gastado
                    reembolso_card = dept_card_anticipado - dept_card_gastado
                    gasto_total = dept_efec_gastado + dept_card_gastado
                    subtotal_dept = dept_efec_anticipado + dept_card_anticipado
                    
                    # Agregar a la lista de resumen
                    dept_summary_data.append({
                        'nombre': dept,
                        'efec_anticipado': dept_efec_anticipado,
                        'card_anticipado': dept_card_anticipado,
                        'efec_gastado': dept_efec_gastado,
                        'card_gastado': dept_card_gastado,
                        'reembolso_efec': reembolso_efec,
                        'reembolso_card': reembolso_card,
                        'gasto_total': gasto_total,
                        'subtotal': subtotal_dept
                    })
                
                # Escribir datos del resumen por departamento
                for dept_summary in dept_summary_data:
                    # DEPARTAMENTO
                    ws.cell(row=summary_row, column=1, value=dept_summary['nombre'])
                    
                    # Total Efectivo Anticipado
                    ws.cell(row=summary_row, column=2, value=dept_summary['efec_anticipado'])
                    ws.cell(row=summary_row, column=2).number_format = '0.00'
                    
                    # Total Tarjeta Anticipado
                    ws.cell(row=summary_row, column=3, value=dept_summary['card_anticipado'])
                    ws.cell(row=summary_row, column=3).number_format = '0.00'
                    
                    # Total Efectivo Gastado
                    ws.cell(row=summary_row, column=4, value=dept_summary['efec_gastado'])
                    ws.cell(row=summary_row, column=4).number_format = '0.00'
                    
                    # Total Tarjeta Gastado
                    ws.cell(row=summary_row, column=5, value=dept_summary['card_gastado'])
                    ws.cell(row=summary_row, column=5).number_format = '0.00'
                    
                    # Reembolso Efectivo
                    reembolso_efec_cell = ws.cell(row=summary_row, column=6, value=dept_summary['reembolso_efec'])
                    reembolso_efec_cell.number_format = '0.00'
                    # Color: rojo si es negativo (faltante), verde si es positivo (sobrante)
                    if dept_summary['reembolso_efec'] < 0:
                        reembolso_efec_cell.font = Font(color="FF0000")  # Rojo
                    elif dept_summary['reembolso_efec'] > 0:
                        reembolso_efec_cell.font = Font(color="00AA00")  # Verde
                    
                    # Reembolso Tarjeta
                    reembolso_card_cell = ws.cell(row=summary_row, column=7, value=dept_summary['reembolso_card'])
                    reembolso_card_cell.number_format = '0.00'
                    if dept_summary['reembolso_card'] < 0:
                        reembolso_card_cell.font = Font(color="FF0000")
                    elif dept_summary['reembolso_card'] > 0:
                        reembolso_card_cell.font = Font(color="00AA00")
                    
                    # Gasto Total
                    ws.cell(row=summary_row, column=8, value=dept_summary['gasto_total'])
                    ws.cell(row=summary_row, column=8).number_format = '0.00'
                    
                    # Subtotal Departamento
                    ws.cell(row=summary_row, column=9, value=dept_summary['subtotal'])
                    ws.cell(row=summary_row, column=9).number_format = '0.00'
                    ws.cell(row=summary_row, column=9).font = Font(bold=True)
                    
                    # Fondo alternado para filas
                    if (summary_row - 3) % 2 == 0:
                        for col in range(1, len(dept_summary_headers) + 1):
                            ws.cell(row=summary_row, column=col).fill = PatternFill(
                                start_color="F8F9FA", end_color="F8F9FA", fill_type="solid"
                            )
                    
                    summary_row += 1
                
                # Fila de TOTALES POR DEPARTAMENTOS
                total_row = summary_row
                
                # "TOTAL GENERAL" en primera columna
                ws.cell(row=total_row, column=1, value="TOTAL GENERAL").font = Font(bold=True, color="FFFFFF")
                
                # Total Efectivo Anticipado
                total_efec_anticipado_cell = ws.cell(row=total_row, column=2, value=general_efec_anticipado)
                total_efec_anticipado_cell.number_format = '0.00'
                total_efec_anticipado_cell.font = Font(bold=True, color="FFFFFF")
                
                # Total Tarjeta Anticipado
                total_card_anticipado_cell = ws.cell(row=total_row, column=3, value=general_card_anticipado)
                total_card_anticipado_cell.number_format = '0.00'
                total_card_anticipado_cell.font = Font(bold=True, color="FFFFFF")
                
                # Total Efectivo Gastado
                total_efec_gastado_cell = ws.cell(row=total_row, column=4, value=general_efec_gastado)
                total_efec_gastado_cell.number_format = '0.00'
                total_efec_gastado_cell.font = Font(bold=True, color="FFFFFF")
                
                # Total Tarjeta Gastado
                total_card_gastado_cell = ws.cell(row=total_row, column=5, value=general_card_gastado)
                total_card_gastado_cell.number_format = '0.00'
                total_card_gastado_cell.font = Font(bold=True, color="FFFFFF")
                
                # Reembolso Efectivo Total
                reembolso_efec_total = general_efec_anticipado - general_efec_gastado
                reembolso_efec_total_cell = ws.cell(row=total_row, column=6, value=reembolso_efec_total)
                reembolso_efec_total_cell.number_format = '0.00'
                reembolso_efec_total_cell.font = Font(bold=True, color="FFFFFF")
                if reembolso_efec_total < 0:
                    reembolso_efec_total_cell.font = Font(bold=True, color="FFCCCC")  # Rojo claro
                elif reembolso_efec_total > 0:
                    reembolso_efec_total_cell.font = Font(bold=True, color="CCFFCC")  # Verde claro
                
                # Reembolso Tarjeta Total
                reembolso_card_total = general_card_anticipado - general_card_gastado
                reembolso_card_total_cell = ws.cell(row=total_row, column=7, value=reembolso_card_total)
                reembolso_card_total_cell.number_format = '0.00'
                reembolso_card_total_cell.font = Font(bold=True, color="FFFFFF")
                if reembolso_card_total < 0:
                    reembolso_card_total_cell.font = Font(bold=True, color="FFCCCC")
                elif reembolso_card_total > 0:
                    reembolso_card_total_cell.font = Font(bold=True, color="CCFFCC")
                
                # Gasto Total General
                gasto_total_general = general_efec_gastado + general_card_gastado
                gasto_total_cell = ws.cell(row=total_row, column=8, value=gasto_total_general)
                gasto_total_cell.number_format = '0.00'
                gasto_total_cell.font = Font(bold=True, color="FFFFFF")
                
                # Subtotal General (Total Anticipado)
                subtotal_general = general_efec_anticipado + general_card_anticipado
                subtotal_cell = ws.cell(row=total_row, column=9, value=subtotal_general)
                subtotal_cell.number_format = '0.00'
                subtotal_cell.font = Font(bold=True, color="FFFFFF")
                
                # Aplicar fondo verde a toda la fila de total general
                for col in range(1, len(dept_summary_headers) + 1):
                    ws.cell(row=total_row, column=col).fill = PatternFill(
                        start_color="27ae60", end_color="27ae60", fill_type="solid"
                    )
                
                # Espacio despu√©s del resumen
                summary_row = total_row + 2
                
                # --- RESUMEN CONSOLIDADO SIMPLE (opcional) ---
                # T√≠tulo del resumen consolidado
                ws.merge_cells(f'A{summary_row}:{get_column_letter(len(headers))}{summary_row}')
                consol_title = ws.cell(row=summary_row, column=1)
                consol_title.value = "üìä RESUMEN CONSOLIDADO"
                consol_title.font = Font(bold=True, color="2c3e50", size=12)
                consol_title.alignment = Alignment(horizontal='center')
                consol_title.fill = PatternFill(start_color="E8F8F5", end_color="E8F8F5", fill_type="solid")
                
                summary_row += 1
                
                # Encabezados del resumen consolidado
                ws.cell(row=summary_row, column=1, value="CONCEPTO").font = Font(bold=True)
                ws.cell(row=summary_row, column=2, value="MONTO").font = Font(bold=True)
                
                summary_row += 1
                
                # Datos del resumen consolidado
                consolidado_data = [
                    ("Total Efectivo Anticipado", general_efec_anticipado),
                    ("Total Tarjeta Anticipado", general_card_anticipado),
                    ("Total Anticipado", subtotal_general),
                    ("", ""),  # Separador
                    ("Total Efectivo Gastado", general_efec_gastado),
                    ("Total Tarjeta Gastado", general_card_gastado),
                    ("Total Gastado", gasto_total_general),
                    ("", ""),  # Separador
                    ("Reembolso Efectivo", reembolso_efec_total),
                    ("Reembolso Tarjeta", reembolso_card_total),
                    ("Reembolso Total", reembolso_efec_total + reembolso_card_total),
                ]
                
                for idx, (concepto, monto) in enumerate(consolidado_data):
                    row = summary_row + idx
                    if concepto:  # Si no es separador
                        ws.cell(row=row, column=1, value=concepto)
                        if monto != "":  # Si tiene monto
                            cell = ws.cell(row=row, column=2, value=monto)
                            cell.number_format = '0.00'
                            
                            # Formato especial para reembolsos
                            if "Reembolso" in concepto:
                                if monto < 0:
                                    cell.font = Font(color="FF0000", bold=True)
                                elif monto > 0:
                                    cell.font = Font(color="00AA00", bold=True)
                            
                            # Formato para totales
                            if "Total" in concepto and "Anticipado" not in concepto and "Gastado" not in concepto:
                                cell.font = Font(bold=True)
                
                # Ajustar el ancho de las columnas del resumen
                ws.column_dimensions['A'].width = 25
                ws.column_dimensions['B'].width = 15
                
            else:
                # Si no hay columna de departamento, tabla simple
                # Encabezados
                for col_idx, header in enumerate(headers, start=1):
                    header_cell = ws.cell(row=current_row, column=col_idx)
                    header_cell.value = header
                    header_cell.font = Font(bold=True, color="FFFFFF")
                    header_cell.fill = PatternFill(start_color="2c3e50", end_color="2c3e50", fill_type="solid")
                    header_cell.alignment = Alignment(horizontal='center', vertical='center')
                    
                    col_letter = get_column_letter(col_idx)
                    ws.column_dimensions[col_letter].width = max(len(str(header)) + 2, 12)
                
                current_row += 1
                
                # Datos
                for row_idx, row in enumerate(data):
                    for col_idx, cell_data in enumerate(row, start=1):
                        cell = ws.cell(row=current_row, column=col_idx, value=cell_data)
                        
                        # Convertir montos a n√∫meros
                        cell_str = str(cell_data) if cell_data is not None else ""
                        if cell_str.startswith('$'):
                            clean_str = cell_str.replace('$', '').replace(',', '').replace(' ', '').strip()
                            if clean_str.replace('.', '', 1).isdigit():
                                try:
                                    cell.value = float(clean_str)
                                    cell.number_format = '0.00'
                                    cell.alignment = Alignment(horizontal='right')
                                except (ValueError, TypeError):
                                    cell.value = cell_str
                        
                    # Fondo alternado
                    if row_idx % 2 == 0:
                        for col_idx in range(1, len(headers) + 1):
                            ws.cell(row=current_row, column=col_idx).fill = PatternFill(
                                start_color="F8F9FA", end_color="F8F9FA", fill_type="solid"
                            )
                    
                    current_row += 1
            
            # Aplicar bordes a todas las celdas con datos
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            
            max_row = ws.max_row
            max_col = ws.max_column
            
            for row in ws.iter_rows(min_row=4, max_row=max_row, min_col=1, max_col=max_col):
                for cell in row:
                    if cell.value is not None:
                        cell.border = thin_border
            
            # Ajustar autom√°ticamente el ancho de las columnas
            for col_idx in range(1, max_col + 1):
                max_length = 0
                col_letter = get_column_letter(col_idx)
                
                for row in range(4, max_row + 1):
                    cell = ws.cell(row=row, column=col_idx)
                    if cell.value:
                        cell_length = len(str(cell.value))
                        if cell_length > max_length:
                            max_length = cell_length
                
                # Ajustar ancho (m√≠nimo 10, m√°ximo 30)
                adjusted_width = min(max(max_length + 2, 10), 30)
                ws.column_dimensions[col_letter].width = adjusted_width
            
            # Guardar archivo
            wb.save(filename)
            
            # Mensaje de √©xito
            if indices['departamento'] is not None:
                message = (
                    f"‚úÖ EXCEL CON TODAS LAS COLUMNAS:\n\n"
                    f"üìÇ {filename}\n\n"
                )
            else:
                message = f"‚úÖ Excel con todas las columnas exportado:\n\n{filename}"
            
            messagebox.showinfo("Exportaci√≥n exitosa", message)
            return filename
            
        except Exception as e:
            messagebox.showerror("‚ùå Error", f"No se pudo exportar a Excel:\n\n{str(e)}")
            import traceback
            traceback.print_exc()
            return None
        
@staticmethod
def create_export_button(parent, tree: ttk.Treeview, title: str, 
                        button_text: str = "üì§ Exportar", 
                        pack_options: dict = None,
                        include_print: bool = True) -> ttk.Button:
    """
    Crea un bot√≥n de exportaci√≥n con men√∫ desplegable
    
    NOTA: Ahora todas las exportaciones son autom√°ticamente jer√°rquicas
          cuando detectan columnas de departamento y solicitante
    """
    from tkinter import Menu
    
    def show_export_menu(event=None):
        """Muestra el men√∫ de exportaci√≥n"""
        menu = Menu(parent, tearoff=0)
        
        # Obtener informaci√≥n sobre la estructura
        headers, _, hierarchical_structure = TreeviewExporter.get_treeview_data(tree, hierarchical=True)
        
        # Solo agregar opciones disponibles
        if HAS_EXCEL:
            menu.add_command(
                label="üìä Excel (.xlsx)",
                command=lambda: TreeviewExporter.export_to_excel(tree, title),
                font=('Arial', 10)
            )

            menu.add_command(
                label="üìä Excel Completo (Todas columnas)",
                command=lambda: TreeviewExporter.export_to_excel_full_columns(tree, title),
                font=('Arial', 10)
            )

        if HAS_WORD:
            menu.add_command(
                label="üìù Word (.docx)",
                command=lambda: TreeviewExporter.export_to_word(tree, title),
                font=('Arial', 10)
            )
        
        if HAS_PDF:
            menu.add_command(
                label="üìÑ PDF (.pdf)",
                command=lambda: TreeviewExporter.export_to_pdf(tree, title),
                font=('Arial', 10)
            )
        
        menu.add_separator()
        
        if include_print and HAS_PDF:
            menu.add_command(
                label="üñ®Ô∏è Imprimir",
                command=lambda: TreeviewExporter.print_directly(tree, title),
                font=('Arial', 10)
            )
            
        menu.add_separator()
        
        # Mostrar men√∫
        if event:
            try:
                menu.tk_popup(event.x_root, event.y_root)
            finally:
                menu.grab_release()
        else:
            x = parent.winfo_rootx() + export_btn.winfo_x()
            y = parent.winfo_rooty() + export_btn.winfo_y() + export_btn.winfo_height()
            menu.tk_popup(x, y)
    
    # Crear bot√≥n
    export_btn = ttk.Button(
        parent,
        text=button_text,
        command=show_export_menu
    )
    
    # Configurar pack si se proporcionaron opciones
    if pack_options:
        export_btn.pack(**pack_options)
    else:
        export_btn.pack(side=tk.RIGHT, padx=5, pady=5)
    
    # Tambi√©n agregar men√∫ contextual al Treeview
    tree.bind("<Button-3>", show_export_menu)
    
    return export_btn